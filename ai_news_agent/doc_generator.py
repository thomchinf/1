from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .models import Article, SectionBundle


class DocxReportGenerator:
    def __init__(self, config: dict[str, Any], logger) -> None:
        self.config = config
        self.logger = logger
        self.doc_cfg = config["document"]

    def generate(self, output_path: Path, sections: list[SectionBundle], metadata: dict[str, Any]) -> Path:
        document = Document()
        self._set_page_layout(document)
        self._build_cover(document, metadata)

        for section in sections:
            self._add_section(document, section)

        document.save(output_path)
        return output_path

    def _set_page_layout(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.0)

    def _build_cover(self, document: Document, metadata: dict[str, Any]) -> None:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(self.doc_cfg.get("title", "每日AI资讯摘编"))
        self._set_run_style(
            title_run,
            self.doc_cfg.get("title_font", "方正小标宋_GBK"),
            self.doc_cfg.get("title_size_pt", 22),
        )

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(
            f"{self.doc_cfg.get('subtitle_prefix', '内部参考')}  |  生成时间：{metadata['generated_at']}"
        )
        self._set_run_style(subtitle_run, self.doc_cfg.get("body_font", "仿宋_GB2312"), 12)

        summary = document.add_paragraph()
        summary.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        summary.paragraph_format.line_spacing = Pt(self.doc_cfg.get("line_spacing_pt", 28))
        summary_run = summary.add_run(
            f"本期共汇总候选资讯 {metadata['candidate_count']} 条，入选正文 {metadata['article_count']} 条。"
            f"数据范围以近 {metadata['recent_hours']} 小时公开信息为主，模型摘要模式：{metadata['llm_mode']}。"
        )
        self._set_run_style(summary_run, self.doc_cfg.get("body_font", "仿宋_GB2312"), self.doc_cfg.get("body_size_pt", 14))

        divider = document.add_paragraph()
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        divider_run = divider.add_run("-" * 56)
        self._set_run_style(divider_run, self.doc_cfg.get("heading_font", "黑体"), 12, color="666666")

    def _add_section(self, document: Document, section: SectionBundle) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(6)
        heading_run = heading.add_run(section.label)
        self._set_run_style(
            heading_run,
            self.doc_cfg.get("heading_font", "黑体"),
            self.doc_cfg.get("heading_size_pt", 16),
            bold=True,
        )

        self._add_section_images(document, section)

        if not section.articles:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(28)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(self.doc_cfg.get("line_spacing_pt", 28))
            run = paragraph.add_run("暂无相关内容")
            self._set_run_style(run, self.doc_cfg.get("body_font", "仿宋_GB2312"), self.doc_cfg.get("body_size_pt", 14))
            return

        for idx, article in enumerate(section.articles, start=1):
            self._add_article_block(document, idx, article)

    def _add_article_block(self, document: Document, idx: int, article: Article) -> None:
        published_label = (
            article.published_at.strftime("%Y-%m-%d %H:%M")
            if article.published_at
            else str(article.metadata.get("published_text") or "未知")
        )

        # V5.0：标题旁显示主体
        entity = getattr(article, 'entity', '') or getattr(article, 'industry', '')
        title_text = f"{idx}. {article.display_title}"
        if entity:
            title_text += f" [主体: {entity}]"

        title_paragraph = document.add_paragraph()
        title_paragraph.paragraph_format.space_before = Pt(4)
        title_paragraph.paragraph_format.space_after = Pt(2)
        title_run = title_paragraph.add_run(title_text)
        self._set_run_style(title_run, self.doc_cfg.get("heading_font", "黑体"), 14, bold=True)

        meta_paragraph = document.add_paragraph()
        meta_run = meta_paragraph.add_run(
            f"来源：{article.source_name}    发布时间：{published_label}"
        )
        self._set_run_style(meta_run, self.doc_cfg.get("emphasis_font", "楷体_GB2312"), 12, color="666666")

        summary_paragraph = document.add_paragraph()
        summary_paragraph.paragraph_format.first_line_indent = Pt(28)
        summary_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        summary_paragraph.paragraph_format.line_spacing = Pt(self.doc_cfg.get("line_spacing_pt", 28))
        summary_run = summary_paragraph.add_run(f"摘要：{article.summary}")
        self._set_run_style(summary_run, self.doc_cfg.get("body_font", "仿宋_GB2312"), self.doc_cfg.get("body_size_pt", 14))

        if self.doc_cfg.get("include_url", True):
            url_paragraph = document.add_paragraph()
            url_paragraph.paragraph_format.first_line_indent = Pt(28)
            url_run = url_paragraph.add_run(f"原文链接：{article.url}")
            self._set_run_style(url_run, self.doc_cfg.get("body_font", "仿宋_GB2312"), 11, color="555555")

    def _add_section_images(self, document: Document, section: SectionBundle) -> None:
        inserted = 0
        for article in section.articles:
            for image_path in article.local_image_paths:
                if not image_path.exists():
                    continue
                document.add_picture(str(image_path), width=Inches(self.doc_cfg.get("image_width_inches", 5.8)))
                caption = document.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run(f"图{inserted + 1}：{article.display_title}（来源：{article.source_name}）")
                self._set_run_style(
                    caption_run,
                    self.doc_cfg.get("body_font", "仿宋_GB2312"),
                    self.doc_cfg.get("caption_size_pt", 11),
                    color="666666",
                )
                inserted += 1
                if inserted >= 2:
                    return

    def _add_finance_table(self, document: Document, articles: list[Article]) -> None:
        rows = [article.finance_info for article in articles if any(article.finance_info.values())]
        if not rows:
            return

        heading = document.add_paragraph()
        run = heading.add_run("投融资事件汇总表")
        self._set_run_style(run, self.doc_cfg.get("heading_font", "黑体"), 13, bold=True)

        table = document.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["企业/项目", "轮次", "金额", "投资方", "业务方向"]
        for cell, header in zip(table.rows[0].cells, headers):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(header)
            self._set_run_style(run, self.doc_cfg.get("heading_font", "黑体"), 11, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row in rows:
            cells = table.add_row().cells
            values = [
                row.get("company", ""),
                row.get("round", ""),
                row.get("amount", ""),
                row.get("investors", ""),
                row.get("business", ""),
            ]
            for cell, value in zip(cells, values):
                run = cell.paragraphs[0].add_run(value)
                self._set_run_style(run, self.doc_cfg.get("body_font", "仿宋_GB2312"), 10.5)

    def _add_paper_table(self, document: Document, articles: list[Article]) -> None:
        rows = [article.paper_info for article in articles if any(article.paper_info.values())]
        if not rows:
            return

        heading = document.add_paragraph()
        run = heading.add_run("论文要点汇总表")
        self._set_run_style(run, self.doc_cfg.get("heading_font", "黑体"), 13, bold=True)

        table = document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["论文来源/会议信息", "研究机构", "核心结论"]
        for cell, header in zip(table.rows[0].cells, headers):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(header)
            self._set_run_style(run, self.doc_cfg.get("heading_font", "黑体"), 11, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row in rows:
            cells = table.add_row().cells
            values = [
                row.get("venue", ""),
                row.get("institution", ""),
                row.get("takeaway", ""),
            ]
            for cell, value in zip(cells, values):
                run = cell.paragraphs[0].add_run(value)
                self._set_run_style(run, self.doc_cfg.get("body_font", "仿宋_GB2312"), 10.5)

    def _set_run_style(
        self,
        run,
        font_name: str,
        font_size_pt: float,
        *,
        bold: bool = False,
        color: str | None = None,
    ) -> None:
        run.font.name = font_name
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(font_size_pt)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
